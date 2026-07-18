import logging
import os
import tempfile
import time
import uuid
import docx
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.models.srs_version import SRSVersion
from app.models.requirement_atom import RequirementAtom
from app.models.session import Session
from app.services.storage_service import StorageService
from app.services.aria_agent import get_groq_client
from app.utils.prompts import PROMPT_VERSION

logger = logging.getLogger(__name__)
settings = get_settings()

class SRSGenerator:
    @classmethod
    async def generate_srs(cls, session_id: uuid.UUID, db: AsyncSession) -> SRSVersion:
        """
        Generates an SRS Word document (.docx) for the session,
        uploads it to storage, and writes an SRSVersion record to the database.
        """
        start_time = time.time()
        logger.info(f"Starting SRS generation for session: {session_id}")

        # 1. Fetch the session and its requirement atoms
        session_result = await db.execute(select(Session).where(Session.id == session_id))
        session = session_result.scalar_one_or_none()
        if not session:
            raise ValueError(f"Session {session_id} not found.")

        atoms_result = await db.execute(
            select(RequirementAtom)
            .where(RequirementAtom.session_id == session_id)
            .where(RequirementAtom.status == "active")
        )
        atoms = atoms_result.scalars().all()

        # 2. Call Groq to generate a professional project summary/intro (optional but nice)
        summary_text = "No summary generated."
        llm_model = settings.GROQ_MODEL
        if atoms:
            client = get_groq_client()
            if not (settings.GROQ_API_KEY.startswith("test") or settings.GROQ_API_KEY.startswith("mock")):
                try:
                    prompt = (
                        "Write a short, professional executive summary (max 300 words) for a Software Requirements "
                        "Specification (SRS) based on the following requirement statements:\n"
                        + "\n".join([f"- {a.raw_text}" for a in atoms])
                    )
                    response = client.chat.completions.create(
                        model=settings.GROQ_MODEL,
                        messages=[{"role": "user", "content": prompt}],
                        timeout=settings.GROQ_TIMEOUT_SECONDS
                    )
                    summary_text = response.choices[0].message.content.strip()
                except Exception as e:
                    logger.error(f"Failed to generate Groq summary for SRS: {e}")
                    summary_text = "Executive summary generation timed out/failed."
            else:
                summary_text = "Mock executive summary for requirement atoms."

        # 3. Create the Word document using python-docx
        doc = docx.Document()
        doc.add_heading("Software Requirements Specification (SRS)", 0)
        doc.add_heading("1. Executive Summary", level=1)
        doc.add_paragraph(summary_text)

        doc.add_heading("2. Functional Requirements", level=1)
        if not atoms:
            doc.add_paragraph("No requirements captured during this session.")
        else:
            for idx, atom in enumerate(atoms, start=1):
                doc.add_heading(f"2.{idx} Requirement {idx}", level=2)
                p = doc.add_paragraph()
                p.add_run("Subject: ").bold = True
                p.add_run(f"{atom.subject}\n")
                p.add_run("Action: ").bold = True
                p.add_run(f"{atom.action}\n")
                if atom.constraint_text:
                    p.add_run("Constraint: ").bold = True
                    p.add_run(f"{atom.constraint_text}\n")
                p.add_run("Raw statement: ").bold = True
                p.add_run(f'"{atom.raw_text}"')

        # 4. Save to a temporary file
        fd, temp_path = tempfile.mkstemp(suffix=".docx")
        try:
            os.close(fd)
            doc.save(temp_path)

            # 5. Determine version number (e.g. read existing versions for project)
            version_q = select(SRSVersion).where(SRSVersion.project_id == session.project_id)
            version_result = await db.execute(version_q)
            existing_count = len(version_result.scalars().all())
            version_str = f"1.{existing_count}"

            # 6. Upload to S3/R2
            file_url = StorageService.upload_srs(
                local_file_path=temp_path,
                project_id=str(session.project_id),
                version=version_str
            )
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

        generation_latency_ms = int((time.time() - start_time) * 1000)

        # 7. Write SRSVersion row
        srs_version = SRSVersion(
            project_id=session.project_id,
            session_id=session.id,
            version=version_str,
            file_url=file_url,
            generated_by="ARIA",
            change_summary=f"Generated after ending session {session_id}.",
            llm_model=llm_model,
            prompt_version=PROMPT_VERSION,
            generation_latency_ms=generation_latency_ms
        )
        db.add(srs_version)
        await db.commit()
        await db.refresh(srs_version)

        logger.info(f"SRS version {version_str} successfully generated and saved.")
        return srs_version
